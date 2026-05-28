from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config import settings

SECTION_TITLES = (
    ("introduction", "Introduction"),
    ("incident_description", "Incident Description"),
    ("medical_treatment", "Medical Treatment"),
    ("employment_history", "Employment History"),
)


def _background_path(job_id: str) -> Path:
    return Path(settings.JOBS_PATH) / job_id / "background.json"


def _word_path(job_id: str) -> Path:
    return Path(settings.JOBS_PATH) / job_id / "background.docx"


def _load_background(job_id: str) -> dict:
    path = _background_path(job_id)
    if not path.exists():
        raise FileNotFoundError(f"Background not found: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def build_background_docx(job_id: str, *, force: bool = False) -> Path:
    """Build a Word background draft from background.json and persist to the job folder."""
    out_path = _word_path(job_id)
    if out_path.exists() and not force:
        return out_path

    background = _load_background(job_id)
    prepared = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = document.add_heading("Background Section", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = document.add_paragraph()
    meta.add_run(f"Job Reference: {job_id}").bold = True
    meta.add_run(f"\nPrepared: {prepared}")

    document.add_paragraph("")

    wrote_section = False
    for key, heading in SECTION_TITLES:
        text = str(background.get(key, "")).strip()
        if not text:
            continue
        wrote_section = True
        document.add_heading(heading, level=1)
        for paragraph in text.split("\n\n"):
            cleaned = paragraph.strip()
            if cleaned:
                document.add_paragraph(cleaned)

    if not wrote_section:
        full_text = str(background.get("full_text", "")).strip()
        if full_text:
            document.add_heading("Background", level=1)
            for paragraph in full_text.split("\n\n"):
                cleaned = paragraph.strip()
                if cleaned:
                    document.add_paragraph(cleaned)

    if not wrote_section and not str(background.get("full_text", "")).strip():
        document.add_paragraph(
            "No background content is available for this job yet."
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    print(f"[WordExport] saved job_id={job_id} path={out_path}")
    return out_path
